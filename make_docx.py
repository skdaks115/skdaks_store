from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('📂 File Converter Pro 사용자 매뉴얼', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph('본 매뉴얼은 File Converter Pro 프로그램의 주요 기능과 사용 방법, 그리고 가산점 항목인 추가 구현 기능들에 대해 설명합니다.')
    
    # Section 1
    doc.add_heading('1. 📝 프로그램 개요', level=1)
    doc.add_paragraph('File Converter Pro는 대량의 파일을 효율적으로 관리하기 위해 제작된 데스크톱 애플리케이션입니다. customtkinter를 기반으로 한 세련된 다크 그레이(Dark Gray) 테마와 베이비 핑크(Baby Pink) 포인트 색상을 적용하여 시각적 편의성을 높였습니다.')
    
    # Section 2
    doc.add_heading('2. 🚀 주요 기능 가이드', level=1)
    
    doc.add_heading('📁 2.1 폴더 선택 및 파일 목록 표시', level=2)
    doc.add_paragraph('• 사용 방법: 상단의 [📁 폴더 선택] 버튼을 클릭하여 작업할 폴더를 지정합니다.\n'
                      '• 기능: 선택된 폴더 내의 모든 파일이 스크롤 영역에 리스트업됩니다.\n'
                      '• 특징: \'현재 파일명\'과 \'변경 후 파일명\'을 나란히 배치하여 작업 상태를 한눈에 확인할 수 있습니다.')
    
    doc.add_heading('🔄 2.2 확장자 일괄 변경 (메인 기능)', level=2)
    doc.add_paragraph('• 사용 방법:\n'
                      '  1. \'변경할 확장자명\' 입력 칸에 새로운 확장자(예: pdf, txt 등)를 입력합니다.\n'
                      '  2. 입력을 시작하면 \'변경 후 파일명\' 영역에 베이비 핑크색으로 미리보기가 활성화됩니다.\n'
                      '  3. 변경하려는 파일의 체크박스를 선택(혹은 전체 선택)합니다.\n'
                      '  4. [🔄 일괄 변경] 버튼을 누르면 작업이 수행됩니다.\n'
                      '• 안정성: 동일한 이름의 파일이 이미 존재할 경우, 로그 시스템을 통해 경고를 띄우고 작업을 건너뛰어 데이터 손실을 방지합니다.')
    
    doc.add_heading('🔍 2.3 중복 파일 스캔 (가산점 기능)', level=2)
    doc.add_paragraph('• 기능: 단순히 파일명만 비교하는 것이 아니라, 파일의 크기와 MD5 해시(내용물)를 분석하여 물리적으로 동일한 파일을 찾아냅니다.\n'
                      '• 시각화: 중복된 파일들은 리스트에서 연한 핑크색 배경으로 강조되어 쉽게 식별할 수 있습니다.')
    
    doc.add_heading('🗑️ 2.4 파일 일괄 삭제 (가산점 기능)', level=2)
    doc.add_paragraph('• 사용 방법: 삭제할 파일들을 다중 선택한 후 [🗑️ 선택 삭제] 버튼을 클릭합니다.\n'
                      '• 보안: 실수로 인한 삭제를 방지하기 위해 최종 확인 팝업창이 표시됩니다.')
    
    # Section 3
    doc.add_heading('3. 📊 시스템 로그 및 알림', level=1)
    doc.add_paragraph('• 하단 로그창: 모든 작업(성공, 실패, 오류 원인)이 하단 텍스트 영역에 실시간으로 기록됩니다.\n'
                      '• 오류 처리: 문제가 발생할 경우 사용자에게 친숙한 언어로 번역된 팝업창을 띄워 해결 방법을 안내합니다.')
    
    # Section 4
    doc.add_heading('🛠️ 개발자용: 빌드 방법 (PyInstaller)', level=1)
    code = doc.add_paragraph('pyinstaller --noconsole --onefile "FileConverterApp.py"')
    code.runs[0].font.name = 'Courier New'
    
    # Footer
    doc.add_paragraph('\n본 프로그램은 대학 과제 제출용으로 제작되었습니다.').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.save('Manual.docx')
    print("Manual.docx created successfully.")

if __name__ == "__main__":
    create_manual()
